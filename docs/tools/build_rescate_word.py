from __future__ import annotations

from datetime import date
from pathlib import Path
from typing import Iterable, Sequence

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


TOOLS = Path(__file__).resolve().parent
DOCS = TOOLS.parent
GENERATED = DOCS / "generated"
ARTIFACTS = GENERATED / "build_artifacts"
OUT = GENERATED / "Rescate_Academico_Documento_Profesional.docx"

IPN = "6C1D45"
IPN_DARK = "4A1230"
LIGHT = "F7F3F5"
BORDER = "D9D0D5"
TEXT = "1F2933"
MUTED = "5F6772"


def cm(value: float) -> int:
    return int(Cm(value))


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_text(cell, text: str, bold: bool = False, color: str = TEXT, size: int = 9) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    r.bold = bold
    r.font.name = "Arial"
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor.from_string(color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def set_table_width(table, widths: Sequence[int]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))

    grid = tbl.find(qn("w:tblGrid"))
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths[idx]))


def set_table_borders(table, color: str = BORDER, size: str = "6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margins(table, top: int = 90, start: int = 120, bottom: int = 90, end: int = 120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.find(qn("w:tblCellMar"))
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for side, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{side}"))
        if node is None:
            node = OxmlElement(f"w:{side}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def keep_with_next(paragraph) -> None:
    paragraph.paragraph_format.keep_with_next = True


def add_field(paragraph, instr: str, placeholder: str = "") -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    run._r.append(begin)

    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    run._r.append(instr_el)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    run._r.append(sep)
    if placeholder:
        paragraph.add_run(placeholder)

    end_run = paragraph.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_update_fields(doc: Document) -> None:
    settings = doc.settings.element
    update = settings.find(qn("w:updateFields"))
    if update is None:
        update = OxmlElement("w:updateFields")
        settings.append(update)
    update.set(qn("w:val"), "true")


def add_bottom_border(paragraph, color: str = IPN, size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.find(qn("w:pBdr"))
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = borders.find(qn("w:bottom"))
    if bottom is None:
        bottom = OxmlElement("w:bottom")
        borders.append(bottom)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)


def style_run(run, size: int | None = None, bold: bool | None = None, color: str | None = None) -> None:
    run.font.name = "Arial"
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc: Document, text: str = "", style: str | None = None, align=None, bold: bool = False) -> object:
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_after = Pt(6)
    if text:
        r = p.add_run(text)
        style_run(r, bold=bold)
    return p


def add_heading(doc: Document, number: str, title: str, level: int = 1) -> object:
    text = f"{number} {title}" if number else title
    p = doc.add_heading(text, level=level)
    keep_with_next(p)
    return p


def add_bullets(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r)


def add_numbers(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(4)
        r = p.add_run(item)
        style_run(r)


def add_callout(doc: Document, title: str, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    set_table_borders(table, color="E5D9DF")
    set_cell_margins(table, top=140, bottom=140, start=180, end=180)
    cell = table.cell(0, 0)
    set_cell_shading(cell, LIGHT)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    style_run(r, size=10, bold=True, color=IPN)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    r2 = p2.add_run(text)
    style_run(r2, size=10, color=TEXT)
    doc.add_paragraph()


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int], font_size: int = 8) -> object:
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_width(table, widths)
    set_table_borders(table)
    set_cell_margins(table)
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, IPN)
        set_cell_text(cell, h, bold=True, color="FFFFFF", size=font_size)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    table.rows[0]._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            set_cell_text(cells[i], str(text), size=font_size)
            if i == 0 or str(text).strip() in {"Alta", "Media", "Baja", "Pendiente", "Completado", "Parcial", "Sí", "No"}:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    return table


def get_font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont | ImageFont.ImageFont:
    candidates = [
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
        "C:/Windows/Fonts/calibrib.ttf" if bold else "C:/Windows/Fonts/calibri.ttf",
    ]
    for path in candidates:
        if Path(path).exists():
            return ImageFont.truetype(path, size)
    return ImageFont.load_default()


def draw_multiline(draw: ImageDraw.ImageDraw, xy: tuple[int, int], text: str, font, fill: str, width: int, line_gap: int = 4) -> int:
    x, y = xy
    words = text.split()
    line = ""
    for word in words:
        trial = f"{line} {word}".strip()
        if draw.textbbox((0, 0), trial, font=font)[2] <= width:
            line = trial
        else:
            draw.text((x, y), line, font=font, fill=fill)
            y += font.size + line_gap
            line = word
    if line:
        draw.text((x, y), line, font=font, fill=fill)
        y += font.size + line_gap
    return y


def make_logo() -> Path:
    ARTIFACTS.mkdir(exist_ok=True)
    path = ARTIFACTS / "ipn_cecyt13_mark.png"
    img = Image.new("RGBA", (900, 360), (255, 255, 255, 0))
    draw = ImageDraw.Draw(img)
    burgundy = f"#{IPN}"
    draw.rounded_rectangle((22, 22, 878, 338), radius=38, fill="#ffffff", outline=burgundy, width=6)
    draw.ellipse((70, 70, 290, 290), outline=burgundy, width=12)
    draw.text((118, 121), "IPN", font=get_font(58, True), fill=burgundy)
    draw.text((115, 188), "CECyT 13", font=get_font(25, True), fill="#30343B")
    draw.rectangle((330, 82, 346, 278), fill=burgundy)
    draw.text((380, 88), "Instituto Politécnico Nacional", font=get_font(35, True), fill=burgundy)
    draw.text((382, 140), 'CECyT No. 13 "Ricardo Flores Magón"', font=get_font(27), fill="#30343B")
    draw.text((382, 196), "Rescate Académico", font=get_font(42, True), fill="#30343B")
    draw.text((384, 252), "Prevención de deserción estudiantil", font=get_font(22), fill="#5F6772")
    img.save(path)
    return path


def make_er_diagram() -> Path:
    path = ARTIFACTS / "diagrama_er.png"
    img = Image.new("RGB", (1600, 950), "white")
    draw = ImageDraw.Draw(img)
    title = get_font(34, True)
    body = get_font(19)
    head = get_font(22, True)
    draw.text((48, 34), "Diagrama Entidad-Relación - Rescate Académico", font=title, fill=f"#{IPN}")

    entities = {
        "Alumno": (80, 130, ["Matricula PK", "GrupoId FK", "PromedioGlobal", "RiesgoAcademico"]),
        "Calificación": (430, 130, ["Id PK", "AlumnoMatricula FK", "MateriaId FK", "Valor", "Aprobada"]),
        "Materia": (780, 130, ["Id PK", "Clave", "Nombre"]),
        "Grupo": (1130, 130, ["Id PK", "Clave", "Carrera", "ProfesorId FK"]),
        "Tutor": (1130, 390, ["Id PK", "UserId FK", "Especialidad", "EstaActivo"]),
        "Intervención": (80, 390, ["Id PK", "AlumnoMatricula FK", "TutorId FK", "Tipo", "Fecha"]),
        "PlanMejora": (430, 390, ["Id PK", "AlumnoMatricula FK", "TutorId FK", "Estado"]),
        "Predicción": (780, 390, ["Id PK", "AlumnoMatricula FK", "Probabilidad", "NivelRiesgo"]),
        "Convocatoria": (80, 650, ["Id PK", "ProyectoId FK", "CupoMaximo", "PromedioMinimo"]),
        "Postulación": (430, 650, ["Id PK", "AlumnoId FK", "ProyectoId FK", "Estado"]),
        "Proyecto": (780, 650, ["Id PK", "Titulo", "Tipo"]),
        "Notificación / Bitácora": (1130, 650, ["UserId FK", "Acción", "Tabla", "Fecha"]),
    }
    boxes = {}
    for name, (x, y, lines) in entities.items():
        w, h = 300, 185
        boxes[name] = (x, y, x + w, y + h)
        draw.rounded_rectangle((x, y, x + w, y + h), radius=18, fill="#FDFBFC", outline=f"#{IPN}", width=3)
        draw.rectangle((x, y, x + w, y + 45), fill=f"#{IPN}")
        draw.text((x + 16, y + 10), name, font=head, fill="white")
        yy = y + 62
        for line in lines:
            draw.text((x + 18, yy), line, font=body, fill="#27313B")
            yy += 28

    def connect(a: str, b: str, label: str, off: int = 0) -> None:
        ax1, ay1, ax2, ay2 = boxes[a]
        bx1, by1, bx2, by2 = boxes[b]
        start = (ax2, (ay1 + ay2) // 2 + off)
        end = (bx1, (by1 + by2) // 2 + off)
        if ax1 > bx1:
            start = (ax1, (ay1 + ay2) // 2 + off)
            end = (bx2, (by1 + by2) // 2 + off)
        draw.line([start, end], fill="#6B7280", width=3)
        mx, my = (start[0] + end[0]) // 2, (start[1] + end[1]) // 2
        draw.rounded_rectangle((mx - 50, my - 17, mx + 50, my + 17), radius=8, fill="white", outline="#D1D5DB")
        draw.text((mx - 34, my - 12), label, font=get_font(16, True), fill="#4B5563")

    connect("Alumno", "Calificación", "1:N")
    connect("Calificación", "Materia", "N:1")
    connect("Grupo", "Tutor", "N:1")
    connect("Alumno", "Intervención", "1:N")
    connect("Intervención", "PlanMejora", "base")
    connect("PlanMejora", "Predicción", "eval")
    connect("Convocatoria", "Postulación", "1:N")
    connect("Postulación", "Proyecto", "N:1")
    connect("Proyecto", "Notificación / Bitácora", "eventos")
    draw.line([(230, 315), (230, 610), (430, 742)], fill="#6B7280", width=3)
    draw.text((254, 585), "Alumno 1:N Postulación", font=get_font(17, True), fill="#4B5563")
    draw.line([(1280, 315), (1280, 390)], fill="#6B7280", width=3)
    draw.text((1302, 340), "Grupo N:1 Tutor", font=get_font(17, True), fill="#4B5563")
    img.save(path)
    return path


def make_class_diagram() -> Path:
    path = ARTIFACTS / "diagrama_clases.png"
    img = Image.new("RGB", (1600, 1050), "white")
    draw = ImageDraw.Draw(img)
    draw.text((48, 34), "Diagrama de Clases UML - Modelos, Servicios y Controladores", font=get_font(34, True), fill=f"#{IPN}")
    font_h = get_font(21, True)
    font_b = get_font(17)

    columns = [
        ("Controladores MVC", 70, 135, [
            "AccountController",
            "DashboardController",
            "ProfesorController",
            "AlumnosController",
            "ConvocatoriasController",
            "PostulacionesController",
            "ReportesController",
            "AdminController",
        ]),
        ("Servicios de Aplicación", 555, 135, [
            "CurrentUserContext",
            "StudentAccessService",
            "RiskEvaluationService",
            "DesercionPredictionService",
            "ConvocatoriaEligibilityService",
            "NotificationService",
            "FileStorageService",
            "AlertasService",
        ]),
        ("Persistencia y Modelos", 1040, 135, [
            "ApplicationDbContext",
            "IdentityUser / Roles",
            "Alumno, Tutor, Grupo",
            "Materia, Calificacion",
            "Convocatoria, Postulacion",
            "IntervencionTutoria",
            "PlanMejora",
            "PrediccionDesercion",
            "BitacoraLog",
        ]),
    ]
    boxes = []
    for title, x, y, lines in columns:
        w, h = 410, 770
        boxes.append((x, y, x + w, y + h))
        draw.rounded_rectangle((x, y, x + w, y + h), radius=22, fill="#FDFBFC", outline=f"#{IPN}", width=3)
        draw.rectangle((x, y, x + w, y + 54), fill=f"#{IPN}")
        draw.text((x + 22, y + 15), title, font=font_h, fill="white")
        yy = y + 83
        for line in lines:
            draw.rounded_rectangle((x + 28, yy, x + w - 28, yy + 56), radius=12, fill="white", outline="#D8C9D0", width=2)
            draw.text((x + 44, yy + 16), line, font=font_b, fill="#27313B")
            yy += 74

    arrow_color = "#6B7280"
    for a, b, label in [
        (boxes[0], boxes[1], "inyectan y coordinan"),
        (boxes[1], boxes[2], "consultan / persisten"),
    ]:
        ax = a[2]
        ay = (a[1] + a[3]) // 2
        bx = b[0]
        by = (b[1] + b[3]) // 2
        draw.line([(ax + 8, ay), (bx - 8, by)], fill=arrow_color, width=5)
        draw.polygon([(bx - 8, by), (bx - 30, by - 12), (bx - 30, by + 12)], fill=arrow_color)
        mx = (ax + bx) // 2
        draw.rounded_rectangle((mx - 115, ay - 24, mx + 115, ay + 24), radius=12, fill="white", outline="#D1D5DB")
        draw.text((mx - 91, ay - 11), label, font=get_font(16, True), fill="#4B5563")
    draw.rounded_rectangle((130, 940, 1470, 1005), radius=16, fill=f"#{LIGHT}", outline="#D8C9D0")
    draw.text((158, 960), "Dependencias transversales: ASP.NET Core Identity, Entity Framework Core, PostgreSQL/SQLite, OpenAI API y filtros de auditoría.", font=get_font(22), fill="#27313B")
    img.save(path)
    return path


def format_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.85)
    section.right_margin = Inches(0.85)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(10.5)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing = 1.08
    normal.paragraph_format.space_after = Pt(6)

    for name, size, color in [
        ("Title", 24, IPN),
        ("Subtitle", 12, MUTED),
        ("Heading 1", 16, IPN),
        ("Heading 2", 13, IPN_DARK),
        ("Heading 3", 11.5, TEXT),
    ]:
        st = styles[name]
        st.font.name = "Arial"
        st.font.size = Pt(size)
        st.font.bold = True if "Heading" in name or name == "Title" else False
        st.font.color.rgb = RGBColor.from_string(color)
    styles["Heading 1"].paragraph_format.space_before = Pt(14)
    styles["Heading 1"].paragraph_format.space_after = Pt(8)
    styles["Heading 2"].paragraph_format.space_before = Pt(10)
    styles["Heading 2"].paragraph_format.space_after = Pt(5)
    styles["Heading 3"].paragraph_format.space_before = Pt(7)
    styles["Heading 3"].paragraph_format.space_after = Pt(4)


def configure_header_footer(section) -> None:
    header = section.header
    p = header.paragraphs[0]
    p.text = ""
    r = p.add_run("Rescate Académico | IPN CECyT No. 13")
    style_run(r, size=8, color=MUTED)
    add_bottom_border(p, color=IPN, size="6")

    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run("Página ")
    style_run(r, size=8, color=MUTED)
    add_field(p, "PAGE")


def add_cover(doc: Document, logo_path: Path) -> None:
    doc.add_picture(str(logo_path), width=Inches(6.5))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Rescate Académico — Plataforma de Monitoreo de Riesgo Académico")
    style_run(r, size=24, bold=True, color=IPN)
    p.paragraph_format.space_after = Pt(8)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Sistema Web para la Prevención de Deserción Estudiantil")
    style_run(r, size=14, color=MUTED)
    p.paragraph_format.space_after = Pt(24)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Instituto Politécnico Nacional — CECyT No. 13 "Ricardo Flores Magón"')
    style_run(r, size=13, bold=True, color=TEXT)
    p.paragraph_format.space_after = Pt(18)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("12 de mayo de 2026")
    style_run(r, size=11, color=MUTED)
    p.paragraph_format.space_after = Pt(28)

    add_callout(doc, "Equipo de desarrollo", "Proyecto académico desarrollado por cinco integrantes con responsabilidades funcionales y técnicas claramente definidas.")
    rows = [
        ("Sergio", "Seguridad, Roles, Bitácora"),
        ("Sara", "Convocatorias, Postulaciones, Reportes"),
        ("Alejandra", "Dashboard, Perfil Académico, Estadísticas"),
        ("Elías", "Administración, Operación Institucional"),
        ("Buenfil", "Inteligencia Artificial, Predicciones"),
    ]
    add_table(doc, ["Integrante", "Responsabilidad"], rows, [2500, 6860], font_size=10)
    doc.add_page_break()


def add_toc(doc: Document) -> None:
    add_heading(doc, "2.", "Índice (Tabla de Contenido)", 1)
    p = doc.add_paragraph()
    add_field(p, r'TOC \o "1-3" \h \z \u', "Actualice este campo en Word para regenerar el índice automáticamente.")
    p.paragraph_format.space_after = Pt(12)
    add_callout(doc, "Nota de uso", "El índice fue insertado como campo automático de Word basado en los estilos de encabezado del documento. Al abrirlo, seleccione Actualizar tabla si Word solicita refrescar campos.")
    doc.add_page_break()


def add_intro(doc: Document) -> None:
    add_heading(doc, "3.", "Introducción", 1)
    paragraphs = [
        "La deserción escolar es uno de los principales problemas del sistema educativo medio superior en México. En el IPN, los índices de abandono académico representan una preocupación institucional constante. Rescate Académico surge como una solución tecnológica integral para la identificación temprana de estudiantes en riesgo, facilitando intervenciones oportunas por parte del cuerpo docente y las autoridades académicas.",
        "El sistema permite a profesores monitorear el desempeño de sus grupos, crear intervenciones y planes de mejora personalizados, y consultar predicciones de riesgo basadas en inteligencia artificial. Los alumnos pueden consultar su perfil académico completo, simular escenarios de calificaciones y postularse a convocatorias de proyectos académicos. Las autoridades y administradores cuentan con herramientas de análisis estadístico, reportes exportables y gestión integral del padrón estudiantil.",
        "Desarrollado con ASP.NET Core 8 MVC y Entity Framework Core, el sistema se despliega en Railway con PostgreSQL en producción y SQLite en desarrollo local. La arquitectura incorpora una capa de servicios extraídos siguiendo principios de Clean Architecture, con seguridad reforzada mediante CSRF, rate limiting, Content Security Policy y control de acceso centralizado.",
    ]
    for text in paragraphs:
        add_para(doc, text)


def add_scope(doc: Document) -> None:
    add_heading(doc, "4.", "Hoja de Alcance", 1)
    add_para(doc, "El sistema Rescate Académico abarca:")
    add_bullets(doc, [
        "Módulo de autenticación segura con 4 roles: Administrador, Tutor/Profesor, Alumno y Autoridad.",
        "Dashboard dinámico por rol con estadísticas en tiempo real.",
        "Perfil académico integral del alumno con semáforo de riesgo.",
        "Sistema de predicción de deserción con heurística institucional y OpenAI GPT-4o-mini.",
        "Registro de intervenciones de tutoría y planes de mejora.",
        "Catálogo de convocatorias y sistema de postulaciones con validación de elegibilidad.",
        "Subida de documentos con validación MIME y almacenamiento privado.",
        "Sistema de notificaciones con badge en tiempo real.",
        "Exportación de datos en CSV con protección contra inyección.",
        "Reporte institucional imprimible.",
        "Bitácora de auditoría para mutaciones críticas.",
        "Panel de administración con CRUD completo.",
        "Despliegue automatizado en Railway con CI/CD desde GitHub.",
    ])
    add_para(doc, "Queda fuera del alcance:")
    add_bullets(doc, [
        "Aplicación móvil nativa.",
        "Integración con sistemas externos del IPN, como TAXCOM o SAES.",
        "Módulo de reconocimiento facial.",
        "Pasarela de pagos.",
        "Chat en tiempo real.",
    ])


def add_analysis(doc: Document) -> None:
    add_heading(doc, "5.", "Análisis", 1)
    add_heading(doc, "5.1", "Planteamiento del problema", 2)
    add_para(doc, "En las instituciones de educación media superior del IPN, los profesores carecen de herramientas centralizadas para identificar tempranamente a estudiantes en riesgo académico. Actualmente, el seguimiento se realiza de forma manual, con expedientes en papel y comunicación informal, lo que provoca:")
    add_bullets(doc, [
        "Detección tardía de estudiantes con múltiples materias reprobadas.",
        "Falta de trazabilidad en intervenciones de apoyo.",
        "Inexistencia de indicadores predictivos de deserción.",
        "Dificultad para que los alumnos conozcan su situación académica integral.",
        "Procesos de postulación a proyectos académicos desconectados del perfil del estudiante.",
    ])
    add_heading(doc, "5.2", "Contexto de información", 2)
    add_para(doc, 'El proyecto se desarrolla para el CECyT No. 13 "Ricardo Flores Magón" del IPN, una institución de nivel medio superior con carreras técnicas. El sistema está diseñado para operar con datos académicos típicos: promedios, calificaciones por materia, inasistencias, evaluaciones parciales, ETS y recursamientos. La asignación de estudiantes a profesores se realiza mediante grupos académicos.')
    add_heading(doc, "5.3", "Objetivos y oportunidades", 2)
    add_callout(doc, "Objetivo general", "Desarrollar una plataforma web integral para el monitoreo, predicción y atención del riesgo académico estudiantil en el IPN.")
    add_para(doc, "Objetivos específicos:")
    add_numbers(doc, [
        "Implementar un sistema de autenticación segura con roles diferenciados.",
        "Crear un perfil académico integral con indicadores visuales de riesgo.",
        "Desarrollar un motor de predicción de deserción con heurística e IA.",
        "Facilitar la creación de intervenciones y planes de mejora por parte de profesores.",
        "Digitalizar el proceso de postulación a proyectos y convocatorias académicas.",
        "Proveer herramientas de análisis estadístico y exportación de datos para autoridades.",
    ])
    add_para(doc, "Oportunidades:")
    add_bullets(doc, [
        "Reducción del índice de deserción mediante identificación temprana.",
        "Mejora en la comunicación profesor-alumno-autoridad.",
        "Toma de decisiones basada en datos.",
        "Digitalización de procesos administrativos académicos.",
    ])
    add_heading(doc, "5.4", "Estudio de factibilidad", 2)
    feasibility = [
        ("Factibilidad Técnica", "Stack tecnológico maduro y gratuito (.NET 8, PostgreSQL, SQLite), infraestructura administrada en Railway, OpenAI API para análisis narrativo, bibliotecas frontend gratuitas y dominio técnico del equipo. Conclusión: TOTALMENTE FACTIBLE."),
        ("Factibilidad Económica", "Herramientas de desarrollo gratuitas, Railway plan hobby aproximado de $5 USD mensuales, OpenAI API con costo estimado de $0.09 USD por análisis completo de 300 estudiantes y ausencia de licenciamiento. Conclusión: FACTIBLE, costo operativo menor a $10 USD mensuales."),
        ("Factibilidad Operativa", "El sistema se integra al flujo existente de profesores y autoridades mediante interfaz responsiva, modo oscuro, breadcrumbs y notificaciones automáticas. Conclusión: FACTIBLE, curva de aprendizaje baja."),
        ("Factibilidad Legal y Ética", "Los datos académicos se protegen conforme a la Ley Federal de Protección de Datos Personales, con autenticación segura, cookies protegidas, auditoría de mutaciones críticas y consentimiento informado para análisis con IA. Conclusión: FACTIBLE."),
        ("Factibilidad de Tiempo", "Desarrollo completado en aproximadamente 9 semanas: autenticación y modelos base, dashboard y perfil académico, predicción IA e intervenciones, seguridad, pruebas y despliegue. Conclusión: FACTIBLE, proyecto entregado dentro del semestre."),
    ]
    for title, text in feasibility:
        add_heading(doc, f"5.4.{feasibility.index((title, text)) + 1}", title, 3)
        add_para(doc, text)
    add_heading(doc, "5.5", "Informe del resultado del estudio", 2)
    add_para(doc, "El proyecto es viable en todas las dimensiones evaluadas. Se recomienda su implementación y despliegue inmediato, con seguimiento post-implementación para incorporar retroalimentación de usuarios reales.")


def add_backlog(doc: Document) -> None:
    add_heading(doc, "6.", "Product Backlog - Listado General de Historias de Usuario", 1)
    rows = [
        ("HU-RA-01", "Inicio de Sesión Seguro", "Alta", "Completado"),
        ("HU-RA-02", "Recuperación de Contraseña", "Alta", "Completado"),
        ("HU-RA-03", "Dashboard por Rol", "Alta", "Completado"),
        ("HU-RA-04", "Gestión de Roles y Privilegios", "Alta", "Completado"),
        ("HU-RA-05", "Prevención de Intrusiones (Bloqueo)", "Alta", "Completado"),
        ("HU-RA-06", "Consulta de Situación Académica", "Alta", "Completado"),
        ("HU-RA-07", "Catálogo de Convocatorias", "Media", "Completado"),
        ("HU-RA-08", "Postulación a Proyectos", "Alta", "Completado"),
        ("HU-RA-09", "Notificaciones de Acción", "Media", "Completado"),
        ("HU-RA-10", "Monitoreo de Tutorados", "Alta", "Completado"),
        ("HU-RA-11", "Semáforo de Rendimiento", "Alta", "Completado"),
        ("HU-RA-12", "Carga Masiva de Alumnado", "Baja", "Pendiente"),
        ("HU-RA-13", "Ajustes de Ciclos Escolares", "Media", "Completado"),
        ("HU-RA-14", "Bitácora de Auditoría", "Media", "Completado"),
        ("HU-RA-15", "Tablero Estadístico", "Media", "Completado"),
        ("HU-RA-16", "Emisión de Reportes Oficiales", "Media", "Parcial"),
        ("HU-RA-17", "Sugerencia Optimizada (IA)", "Alta", "Completado"),
        ("HU-RA-18", "Prevención de Deserción (IA)", "Alta", "Completado"),
        ("HU-RA-19", "Simulador de Elegibilidad (What-If)", "Media", "Completado"),
        ("HU-RA-20", "Intervención Temprana", "Alta", "Completado"),
        ("HU-RA-21", "Seguimiento de Alertas", "Alta", "Completado"),
        ("HU-RA-22", "Evidencia de Postulación (Documentos)", "Media", "Completado"),
        ("HU-RA-23", "Evaluación Masiva de Riesgo", "Media", "Completado"),
    ]
    add_table(doc, ["ID", "Título", "Prioridad", "Estado"], rows, [1250, 5140, 1380, 1590], font_size=8)


HU_DATA = [
    {
        "id": "HU-RA-01",
        "title": "Inicio de Sesión Seguro",
        "desc": "Como usuario del sistema, quiero iniciar sesión con mi correo institucional para acceder a las funciones según mi rol.",
        "analysis": "El sistema valida credenciales contra Identity, verifica que la cuenta esté activa y no pendiente de verificación, y bloquea tras 3 intentos fallidos por 20 minutos. Tras autenticación exitosa, el profesor es redirigido a Mis Grupos y los demás roles al Dashboard.",
        "diagram": "Login -> Validar credenciales -> Verificar estado -> Bloqueo/Acceso -> Redirigir por rol.",
        "code": "AccountController.Login POST; validación de correo institucional (@ipn.mx, @alumno.ipn.mx); PasswordSignInAsync con lockoutOnFailure: true.",
        "input": "email (string), password (string), rememberMe (bool)",
        "output": "Redirección a Dashboard/Profesor o mensaje de error con conteo de intentos fallidos.",
        "db": "AspNetUsers, AspNetRoles, AspNetUserRoles",
        "criteria": [
            "Usuario con credenciales correctas y cuenta activa accede exitosamente.",
            "Usuario con credenciales incorrectas recibe mensaje de error.",
            "Tras 3 intentos fallidos, la cuenta se bloquea por 20 minutos.",
            "Usuario con cuenta inactiva o pendiente recibe mensaje apropiado.",
            "Profesor es redirigido a Profesor/Index; otros roles a Dashboard/Index.",
        ],
        "tests": [
            ("Login válido admin", "admin@ipn.mx / Admin123!", "Redirige a Dashboard", "Completado"),
            ("Login válido profesor", "tutor@ipn.mx / Tutor123!", "Redirige a Mis Grupos", "Completado"),
            ("Login válido alumno", "alumno@alumno.ipn.mx / Alumno123!", "Redirige a Dashboard", "Completado"),
            ("Contraseña incorrecta", "admin@ipn.mx / wrong", "Error y conteo de intentos", "Completado"),
            ("3 intentos fallidos", "3 contraseñas erróneas", "Bloqueo por 20 minutos", "Completado"),
            ("Cuenta inactiva", "usuario inactivo", "Cuenta no activa", "Completado"),
        ],
        "evidence": "Captura de login exitoso y captura de mensaje de bloqueo.",
    },
    {
        "id": "HU-RA-06",
        "title": "Consulta de Situación Académica",
        "desc": "Como alumno, quiero consultar mi perfil académico integral para conocer mi promedio, materias, riesgo y recomendaciones.",
        "analysis": "El perfil académico consolida calificaciones, materias aprobadas y reprobadas, ETS, recursamientos, promedio global y riesgo académico. La vista se ajusta al alumno autenticado y evita que consulte expedientes ajenos.",
        "diagram": "Alumno autenticado -> PerfilAcadémico/MiPerfil -> Cargar expediente -> Calcular indicadores -> Mostrar historial y sugerencias.",
        "code": "PerfilAcademicoController.MiPerfil; RiskEvaluationService; consultas EF Core sobre Alumno, Calificaciones y Materias.",
        "input": "Usuario autenticado; matrícula asociada al usuario.",
        "output": "Vista de perfil académico con historial, semáforo, estadísticas y sugerencias.",
        "db": "Alumnos, Calificaciones, Materias, PrediccionesDesercion",
        "criteria": [
            "El alumno visualiza únicamente su propio expediente.",
            "El promedio global coincide con los datos de calificaciones.",
            "El semáforo se muestra con color y etiqueta coherente.",
            "La tabla de materias permite identificar aprobadas y reprobadas.",
            "El simulador What-If recalcula escenarios sin persistir cambios.",
        ],
        "tests": [
            ("Perfil propio", "alumno autenticado", "Muestra expediente correcto", "Completado"),
            ("Acceso cruzado", "matrícula ajena", "Acceso denegado", "Completado"),
            ("Promedio", "calificaciones demo", "Promedio calculado", "Completado"),
            ("Materias reprobadas", "calificación < 6", "Marcada como reprobada", "Completado"),
            ("Simulador", "cambio hipotético", "Actualiza riesgo en pantalla", "Completado"),
        ],
        "evidence": "Captura de perfil con semáforo, historial y simulador.",
    },
    {
        "id": "HU-RA-08",
        "title": "Postulación a Proyectos",
        "desc": "Como alumno, quiero postularme a convocatorias académicas disponibles para participar en proyectos acordes a mi perfil.",
        "analysis": "El sistema valida promedio mínimo, semestre, carrera, cupo disponible, estado de la convocatoria y duplicidad de postulación. La misma lógica se aplica en GET y POST mediante ConvocatoriaEligibilityService.",
        "diagram": "Catálogo -> Ver convocatoria -> Evaluar elegibilidad -> Adjuntar documento -> Registrar postulación -> Notificar.",
        "code": "PostulacionesController; ConvocatoriaEligibilityService; FileStorageService; NotificationService.",
        "input": "convocatoriaId, documento opcional, usuario alumno.",
        "output": "Postulación registrada o mensaje de inelegibilidad.",
        "db": "Convocatorias, Proyectos, Postulaciones, Alumnos, Notificaciones",
        "criteria": [
            "Alumno elegible puede postularse exitosamente.",
            "Alumno inelegible recibe causas específicas.",
            "No se permite duplicar postulación a la misma convocatoria.",
            "El cupo máximo se respeta en concurrencia normal.",
            "El documento se almacena con nombre seguro si cumple validaciones.",
        ],
        "tests": [
            ("Elegible", "promedio y semestre válidos", "Postulación creada", "Completado"),
            ("Promedio insuficiente", "promedio menor al mínimo", "Error explicativo", "Completado"),
            ("Cupo lleno", "sin lugares disponibles", "No permite postular", "Completado"),
            ("Documento válido", "PDF permitido", "Archivo guardado", "Completado"),
            ("Documento inválido", "EXE renombrado", "Rechazo MIME", "Completado"),
        ],
        "evidence": "Captura de postulación exitosa y validación de elegibilidad.",
    },
    {
        "id": "HU-RA-10",
        "title": "Monitoreo de Tutorados",
        "desc": "Como profesor, quiero visualizar mis grupos y tutorados para identificar riesgos académicos y dar seguimiento.",
        "analysis": "El profesor accede a grupos asignados mediante Grupo.ProfesorId. StudentAccessService centraliza la autorización para evitar IDOR y filtra la visibilidad de alumnos en todas las consultas.",
        "diagram": "Profesor -> Mis Grupos -> Filtrar por Grupo.ProfesorId -> Mostrar riesgo agregado -> Ver tutorados autorizados.",
        "code": "ProfesorController.Index; AlumnosController.MisTutorados; StudentAccessService.ApplyVisibleStudents().",
        "input": "Usuario profesor autenticado; grupos asignados.",
        "output": "Tarjetas de grupo, distribución de riesgo y tabla de estudiantes visibles.",
        "db": "Tutores, Grupos, Alumnos, Calificaciones, PrediccionesDesercion",
        "criteria": [
            "Profesor ve únicamente grupos asignados.",
            "La distribución de riesgo coincide con alumnos del grupo.",
            "La vista rápida no permite consultar estudiantes de otro grupo.",
            "Los enlaces contextuales respetan permisos.",
            "La página carga con datos demo de 12 profesores y grupos.",
        ],
        "tests": [
            ("Mis grupos", "profesor válido", "Muestra grupos asignados", "Completado"),
            ("Grupo ajeno", "id de otro profesor", "403 o sin datos", "Completado"),
            ("Conteo riesgo", "grupo demo", "Totales correctos", "Completado"),
            ("Vista rápida", "alumno del grupo", "Modal abre", "Completado"),
            ("Vista rápida ajena", "matrícula ajena", "Acceso denegado", "Completado"),
        ],
        "evidence": "Captura de Mis Grupos y modal de vista rápida.",
    },
    {
        "id": "HU-RA-11",
        "title": "Semáforo de Rendimiento",
        "desc": "Como profesor o alumno, quiero ver un semáforo de rendimiento para interpretar rápidamente el nivel de riesgo académico.",
        "analysis": "RiskEvaluationService clasifica el riesgo con base en promedio, reprobaciones, ausencias, ETS y recursamientos. El semáforo usa etiquetas institucionales: Verde, Amarillo y Rojo.",
        "diagram": "Indicadores académicos -> Calcular riesgo -> Asignar nivel -> Mostrar badge, color y sugerencias.",
        "code": "RiskEvaluationService.CalcularRiesgo(); vistas PerfilAcadémico, Profesor y Dashboard.",
        "input": "Promedio, número de materias reprobadas, ausencias, ETS, recursamientos.",
        "output": "Nivel de riesgo, color visual, probabilidad estimada y recomendaciones.",
        "db": "Alumnos, Calificaciones, PrediccionesDesercion",
        "criteria": [
            "Riesgo verde se muestra para desempeño satisfactorio.",
            "Riesgo amarillo se muestra para señales de alerta moderada.",
            "Riesgo rojo se muestra para riesgo alto.",
            "El color se adapta al modo claro y oscuro.",
            "Las sugerencias corresponden al nivel calculado.",
        ],
        "tests": [
            ("Verde", "promedio alto, sin reprobadas", "Etiqueta Verde", "Completado"),
            ("Amarillo", "promedio medio o pocas reprobadas", "Etiqueta Amarillo", "Completado"),
            ("Rojo", "múltiples reprobadas", "Etiqueta Rojo", "Completado"),
            ("Dark mode", "tema oscuro", "Contraste correcto", "Completado"),
            ("Sugerencias", "riesgo alto", "Recomienda intervención", "Completado"),
        ],
        "evidence": "Capturas de semáforo en perfil y dashboard.",
    },
    {
        "id": "HU-RA-18",
        "title": "Prevención de Deserción (IA)",
        "desc": "Como autoridad o profesor, quiero consultar predicciones de deserción para priorizar intervenciones.",
        "analysis": "El sistema combina una heurística institucional auditable con análisis narrativo de OpenAI GPT-4o-mini. Si no existe API key, la heurística permanece disponible y la interfaz informa el motivo.",
        "diagram": "Expediente académico -> Heurística institucional -> Probabilidad -> OpenAI opcional -> Análisis narrativo -> Guardar predicción.",
        "code": "DesercionPredictionService; RiskEvaluationService; PrediccionesDesercionController; [AuditLog].",
        "input": "Matrícula del alumno, expediente académico, OPENAI_API_KEY opcional.",
        "output": "Probabilidad, nivel de riesgo, factores detectados y recomendaciones narrativas.",
        "db": "PrediccionesDesercion, Alumnos, Calificaciones, BitacoraLog",
        "criteria": [
            "La heurística funciona sin API key.",
            "La respuesta IA se estructura en secciones parseables.",
            "El análisis se registra en bitácora.",
            "La tasa de solicitudes se limita para proteger costos.",
            "El profesor solo analiza alumnos autorizados.",
        ],
        "tests": [
            ("Heurística", "sin API key", "Predicción disponible", "Completado"),
            ("OpenAI", "API key válida", "Análisis narrativo", "Completado"),
            ("Rate limit", "más de 5/min", "Solicitud limitada", "Completado"),
            ("Auditoría", "guardar análisis", "Bitácora creada", "Completado"),
            ("IDOR", "alumno ajeno", "Acceso denegado", "Completado"),
        ],
        "evidence": "Captura de análisis IA y registro de predicción.",
    },
    {
        "id": "HU-RA-20",
        "title": "Intervención Temprana",
        "desc": "Como profesor, quiero registrar intervenciones y planes de mejora para apoyar a estudiantes en riesgo.",
        "analysis": "Las intervenciones registran tipo, fecha, descripción y seguimiento. Los planes de mejora definen metas, recomendaciones y estado, generando notificaciones al alumno.",
        "diagram": "Detectar alumno en riesgo -> Crear intervención -> Definir plan -> Notificar alumno -> Seguimiento.",
        "code": "IntervencionesController; PlanesMejoraController; NotificationService; StudentAccessService.",
        "input": "Matrícula, tipo de intervención, descripción, metas, recomendaciones.",
        "output": "Intervención o plan creado con notificación asociada.",
        "db": "IntervencionesTutoria, PlanesMejora, Notificaciones, Alumnos, Tutores",
        "criteria": [
            "Profesor solo crea planes para alumnos de sus grupos.",
            "El sistema valida campos obligatorios.",
            "El alumno recibe notificación del plan.",
            "El estado del plan puede actualizarse.",
            "La intervención queda disponible en el historial.",
        ],
        "tests": [
            ("Crear intervención", "datos válidos", "Registro creado", "Completado"),
            ("Plan de mejora", "metas y recomendaciones", "Plan creado", "Completado"),
            ("Alumno ajeno", "matrícula externa", "Acceso denegado", "Completado"),
            ("Notificación", "plan nuevo", "Badge incrementa", "Completado"),
            ("Seguimiento", "cambio de estado", "Historial actualizado", "Completado"),
        ],
        "evidence": "Captura de plan creado y notificación al alumno.",
    },
    {
        "id": "HU-RA-23",
        "title": "Evaluación Masiva de Riesgo",
        "desc": "Como administrador o autoridad, quiero ejecutar una evaluación masiva para actualizar los niveles de riesgo de todo el alumnado.",
        "analysis": "AlertasService recorre alumnos, recalcula riesgo y genera notificaciones cuando existe cambio significativo. La operación permite mantener el padrón actualizado después de cargas o cambios de calificaciones.",
        "diagram": "Administrador -> Evaluación masiva -> Recalcular cada alumno -> Persistir cambios -> Notificar alertas.",
        "code": "AlertasService; Admin/Operaciones; RiskEvaluationService; NotificationService.",
        "input": "Solicitud POST autenticada con antiforgery token.",
        "output": "Resumen de alumnos evaluados, cambios detectados y alertas generadas.",
        "db": "Alumnos, Calificaciones, Notificaciones, BitacoraLog",
        "criteria": [
            "Solo roles autorizados ejecutan la operación.",
            "La operación requiere token antiforgery.",
            "Los riesgos nulos se corrigen automáticamente.",
            "Se notifican cambios relevantes de riesgo.",
            "La operación queda registrada en bitácora.",
        ],
        "tests": [
            ("Ejecutar evaluación", "admin válido", "Resumen generado", "Completado"),
            ("Sin token CSRF", "POST directo", "Solicitud rechazada", "Completado"),
            ("Riesgo nulo", "alumno sin riesgo", "Riesgo calculado", "Completado"),
            ("Cambio relevante", "sube a rojo", "Notificación creada", "Completado"),
            ("Usuario no autorizado", "alumno", "Acceso denegado", "Completado"),
        ],
        "evidence": "Captura de operación masiva y bitácora.",
    },
]


def add_user_stories(doc: Document) -> None:
    add_heading(doc, "7.", "Detalles de Historias de Usuario", 1)
    for idx, hu in enumerate(HU_DATA, start=1):
        add_heading(doc, f"7.{idx}", f"{hu['id']}: {hu['title']}", 2)
        labels = [
            ("Descripción", hu["desc"]),
            ("Análisis", hu["analysis"]),
            ("Diagrama de solución", hu["diagram"]),
            ("Código de solución", hu["code"]),
            ("Parámetros de entrada", hu["input"]),
            ("Parámetros de salida", hu["output"]),
            ("Componentes de BD", hu["db"]),
        ]
        for label, text in labels:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            r = p.add_run(f"{label}: ")
            style_run(r, bold=True, color=IPN)
            r = p.add_run(text)
            style_run(r)
        add_para(doc, "Criterios de aceptación:", bold=True)
        add_numbers(doc, hu["criteria"])
        add_table(doc, ["Caso", "Entrada", "Resultado esperado", "Estado"], hu["tests"], [2100, 2500, 3380, 1380], font_size=7)
        p = doc.add_paragraph()
        r = p.add_run("Evidencias de prueba: ")
        style_run(r, bold=True, color=IPN)
        r = p.add_run(hu["evidence"])
        style_run(r)


def add_design(doc: Document, er_path: Path, class_path: Path) -> None:
    add_heading(doc, "8.", "Diseño", 1)
    add_heading(doc, "8.1", "Diagrama Entidad-Relación", 2)
    add_para(doc, "Tablas principales y relaciones del modelo de datos institucional.")
    doc.add_picture(str(er_path), width=Inches(6.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figura 1. Diagrama Entidad-Relación de entidades principales.", align=WD_ALIGN_PARAGRAPH.CENTER)
    rows = [
        ("Alumno", "Matricula PK, Nombre, Apellidos, Carrera, SemestreActual, PromedioGlobal, RiesgoAcademico, GrupoId FK"),
        ("Calificacion", "Id PK, AlumnoMatricula FK, MateriaId FK, Periodo, Valor, Aprobada, Tipo, VecesCursada"),
        ("Materia", "Id PK, Clave, Nombre"),
        ("Grupo", "Id PK, Clave, Carrera, Semestre, Turno, ProfesorId FK"),
        ("Tutor", "Id PK, Nombre, Apellidos, Especialidad, UserId FK, EstaActivo"),
        ("Convocatoria", "Id PK, Titulo, ProyectoId FK, CupoMaximo, FechaCierre, PromedioMinimo, SemestreMinimo, CarreraRequerida"),
        ("Postulacion", "Id PK, AlumnoId FK, ProyectoId FK, FechaSolicitud, Estado, DocumentoNombre, DocumentoRuta"),
        ("IntervencionTutoria", "Id PK, AlumnoMatricula FK, TutorId FK, Tipo, Descripcion, Fecha, RequiereSeguimiento"),
        ("PlanMejora", "Id PK, AlumnoMatricula FK, TutorId FK, Recomendaciones, Metas, Estado, FechaCreacion"),
        ("PrediccionDesercion", "Id PK, AlumnoMatricula FK, ProbabilidadDesercion, NivelRiesgo, FactoresDetectados, FechaPrediccion"),
        ("Notificacion / BitacoraLog", "UserId, Titulo/Mensaje o Accion/Tabla/Detalles, Leida, FechaCreacion"),
    ]
    add_table(doc, ["Tabla", "Campos principales"], rows, [2300, 7060], font_size=7)
    add_heading(doc, "8.2", "Relaciones principales", 2)
    add_bullets(doc, [
        "Alumno 1:N Calificacion; Alumno N:1 Grupo; Grupo N:1 Tutor.",
        "Materia 1:N Calificacion; Proyecto 1:N Convocatoria; Convocatoria 1:N Postulacion.",
        "Alumno 1:N Postulacion, IntervencionTutoria, PlanMejora y PrediccionDesercion.",
        "Tutor 1:N IntervencionTutoria y PlanMejora.",
    ])
    add_heading(doc, "8.3", "Diagrama de Clases", 2)
    doc.add_picture(str(class_path), width=Inches(6.9))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_para(doc, "Figura 2. Diagrama UML de modelos, servicios y controladores.", align=WD_ALIGN_PARAGRAPH.CENTER)


def add_development(doc: Document) -> None:
    add_heading(doc, "9.", "Desarrollo", 1)
    add_heading(doc, "9.1", "Stack Tecnológico", 2)
    rows = [
        ("Backend", "ASP.NET Core 8 MVC (.NET 8.0)"),
        ("ORM", "Entity Framework Core 8.0"),
        ("Base de datos local", "SQLite"),
        ("Base de datos producción", "PostgreSQL en Railway"),
        ("Autenticación", "ASP.NET Core Identity"),
        ("Frontend", "Bootstrap 5, DataTables 1.13, GSAP 3.12, CountUp.js, Chart.js"),
        ("IA", "OpenAI GPT-4o-mini + heurística institucional"),
        ("Despliegue", "Railway con Docker y GitHub auto-deploy"),
        ("Control de versiones", "Git + GitHub"),
        ("IDE", "JetBrains Rider"),
    ]
    add_table(doc, ["Capa", "Tecnología"], rows, [2600, 6760], font_size=8)
    add_heading(doc, "9.2", "Servicios Extraídos (Clean Architecture)", 2)
    add_bullets(doc, [
        "CurrentUserContext: encapsula acceso a ClaimsPrincipal desde HttpContext.",
        "StudentAccessService: control de acceso centralizado mediante ApplyVisibleStudents(), CanAccessAlumnoAsync() y GetVisibleMatriculasAsync().",
        "RiskEvaluationService: lógica unificada de riesgo, probabilidad, factores y sugerencias.",
        "ConvocatoriaEligibilityService: validación de elegibilidad para postulaciones con patrón EligibilityResult.",
        "NotificationService: creación de notificaciones con patrón unit-of-work.",
        "FileStorageService: almacenamiento privado con validación MIME y extensión.",
        "AlertasService: evaluación masiva y notificación de cambios de riesgo.",
        "DesercionPredictionService: heurística de predicción e integración OpenAI GPT-4o-mini.",
    ])
    add_heading(doc, "9.3", "Seguridad Implementada", 2)
    add_bullets(doc, [
        "AntiForgery Token global mediante AutoValidateAntiforgeryTokenAttribute.",
        "Rate limiting: login 10/min, registro 3/5min, postulación 5/min y OpenAI 5/min.",
        "Content Security Policy con orígenes permitidos para Google Fonts, CDNJS, jsDelivr y DataTables CDN.",
        "Headers HSTS, X-Frame-Options: DENY, X-Content-Type-Options: nosniff, Referrer-Policy y Permissions-Policy.",
        "Cookies HttpOnly, Secure y SameSite=Strict.",
        "Protección contra Open Redirect con Url.IsLocalUrl.",
        "Protección IDOR centralizada con StudentAccessService.",
        "Auditoría mediante [AuditLog] en mutaciones críticas.",
        "Prevención XSS con textContent en modales AJAX y JsonSerializer.Serialize en Chart.js.",
        "Protección contra CSV injection.",
    ])
    add_heading(doc, "9.4", "Queries de BD Relevantes", 2)
    queries = [
        ("Estudiantes en riesgo por grupo", "SELECT g.Clave, COUNT(*) AS Total, SUM(CASE WHEN a.RiesgoAcademico = 'Rojo' THEN 1 ELSE 0 END) AS EnRojo, SUM(CASE WHEN a.RiesgoAcademico = 'Amarillo' THEN 1 ELSE 0 END) AS EnAmarillo FROM Alumnos a JOIN Grupos g ON a.GrupoId = g.Id WHERE g.ProfesorId = @profesorId GROUP BY g.Clave;"),
        ("Materias más reprobadas por grupo", "SELECT m.Nombre, COUNT(*) AS Reprobadas FROM Calificaciones c JOIN Materias m ON c.MateriaId = m.Id JOIN Alumnos a ON c.AlumnoMatricula = a.Matricula WHERE a.GrupoId = @grupoId AND c.Aprobada = 0 GROUP BY m.Nombre ORDER BY Reprobadas DESC;"),
        ("Tendencia de riesgo por período", "SELECT c.Periodo, AVG(a.PromedioGlobal) AS PromedioGeneral, SUM(CASE WHEN a.RiesgoAcademico = 'Rojo' THEN 1 ELSE 0 END) AS Rojos FROM Alumnos a JOIN Calificaciones c ON a.Matricula = c.AlumnoMatricula GROUP BY c.Periodo ORDER BY c.Periodo;"),
    ]
    add_table(doc, ["Consulta", "SQL"], queries, [2500, 6860], font_size=7)
    add_heading(doc, "9.5", "Código Completo", 2)
    add_para(doc, "El código fuente completo, con aproximadamente 16,000 líneas, está disponible en: https://github.com/Xerea/RescateAcademico. El sistema compila con 0 warnings y 0 errores.")


def add_youtube(doc: Document) -> None:
    add_heading(doc, "10.", "Comercial / Link de YouTube", 1)
    add_callout(doc, "Video demostrativo", "Enlace pendiente de integración: se propone un recorrido de 5 a 8 minutos mostrando login de los cuatro roles, dashboard, perfil académico, predicción IA, creación de plan de mejora y postulación a convocatoria.")
    add_para(doc, "URL sugerida para registro final: https://www.youtube.com/ (reemplazar por el enlace definitivo del video).")


def add_general_tests(doc: Document) -> None:
    add_heading(doc, "11.", "Matriz de Pruebas General", 1)
    rows = [
        ("P01", "Login", "Iniciar sesión con credenciales válidas", "Acceso al sistema", "Completado"),
        ("P02", "Login", "Contraseña incorrecta", "Error + conteo de intentos", "Completado"),
        ("P03", "Login", "3 intentos fallidos", "Bloqueo 20 min", "Completado"),
        ("P04", "Roles", "Login como Admin", "Acceso a Gestión y Reportes", "Completado"),
        ("P05", "Roles", "Login como Tutor", "Redirige a Mis Grupos", "Completado"),
        ("P06", "Roles", "Login como Alumno", "Perfil académico y Convocatorias", "Completado"),
        ("P07", "Roles", "Login como Autoridad", "Alumnos, Reportes y Estadísticas", "Completado"),
        ("P08", "Dashboard", "Carga de estadísticas", "Tarjetas con conteos reales", "Completado"),
        ("P09", "Perfil", "Ver perfil académico", "Promedio, semáforo y materias", "Completado"),
        ("P10", "IA", "Predicción de deserción", "Porcentaje y nivel de riesgo", "Completado"),
        ("P11", "IA", "Análisis OpenAI", "Respuesta estructurada con tipeo", "Completado*"),
        ("P12", "IA", "Sin API key configurada", "Mensaje informativo sin error", "Completado"),
        ("P13", "Profesor", "Ver Mis Grupos", "Tarjetas de grupo con riesgo", "Completado"),
        ("P14", "Profesor", "Ver estudiantes de grupo", "Tabla filtrada por grupo", "Completado"),
        ("P15", "Profesor", "Crear plan de mejora", "Plan creado y notificación", "Completado"),
        ("P16", "Profesor", "Tutor bloqueado en plan", "No asigna otro tutor", "Completado"),
        ("P17", "Alumno", "Simulador What-If", "Promedio y riesgo en tiempo real", "Completado"),
        ("P18", "Postulación", "Postularse", "Validación de elegibilidad", "Completado"),
        ("P19", "Postulación", "Cupo excedido", "Mensaje de error", "Completado"),
        ("P20", "Postulación", "Subir documento", "Archivo guardado seguro", "Completado"),
        ("P21", "Postulación", "Archivo no permitido", "Error MIME", "Completado"),
        ("P22", "Export", "CSV de alumnos", "Descarga correcta", "Completado"),
        ("P23", "Seguridad", "CSRF sin token", "Solicitud rechazada", "Completado"),
        ("P24", "Seguridad", "Acceso a alumno de otro grupo", "403", "Completado"),
        ("P25", "UI", "Modo oscuro", "Cambio correcto", "Completado"),
        ("P26", "UI", "DataTables", "Orden, búsqueda y paginación", "Completado"),
        ("P27", "UI", "Quick View modal", "Carga sin congelar pantalla", "Completado"),
    ]
    add_table(doc, ["ID", "Módulo", "Prueba", "Resultado esperado", "Estado"], rows, [800, 1300, 3300, 2800, 1160], font_size=7)
    add_para(doc, "*Requiere OPENAI_API_KEY configurada.")


def add_evidence(doc: Document) -> None:
    add_heading(doc, "12.", "Evidencias de Pruebas General", 1)
    add_para(doc, "Se consideran como evidencias mínimas para la entrega final las siguientes capturas de pantalla del sistema en operación:")
    add_numbers(doc, [
        "Login exitoso Admin.",
        "Login bloqueado tras 3 intentos.",
        "Dashboard Admin con estadísticas.",
        "Mis Grupos con tarjetas y predicciones.",
        "Perfil académico con semáforo y probabilidad.",
        "Predicción IA con análisis narrativo.",
        "Plan de mejora creado.",
        "Postulación exitosa.",
        "CSV exportado.",
        "Modo oscuro activado.",
    ])
    add_callout(doc, "Formato de inserción de evidencias", "Cada captura debe integrarse con título, fecha, usuario de prueba, descripción del resultado y módulo correspondiente. Se recomienda mantener resolución legible y ocultar datos sensibles si se usan alumnos reales.")


def add_user_manual(doc: Document) -> None:
    add_heading(doc, "13.", "Manual de Usuario", 1)
    sections = [
        ("Acceso al Sistema", [
            "Abrir el navegador y acceder a la URL del sistema.",
            "En la página de login, ingresar correo institucional y contraseña.",
            "Si es la primera vez, hacer clic en Registrarse y completar el formulario.",
        ]),
        ("Para Profesores (Tutores)", [
            "Al iniciar sesión, se muestra la página Mis Grupos.",
            "Revisar la distribución de riesgo y las predicciones IA de los estudiantes.",
            "Hacer clic en un grupo para ver la lista de estudiantes.",
            "Usar el ícono de ojo para vista rápida de un estudiante.",
            "Usar el ícono de gráfica para ver el historial académico completo.",
            "Desde el historial, crear un plan de mejora o consultar el análisis IA.",
            "Usar Acciones Rápidas para ver materias reprobadas, intervenciones y planes.",
        ]),
        ("Para Alumnos", [
            "Al iniciar sesión, revisar el panel con resumen de situación académica.",
            "En Mi Perfil Académico, consultar promedio, riesgo y materias.",
            "Usar el simulador ¿Cómo me va? para proyectar escenarios.",
            "En Convocatorias, explorar proyectos disponibles y postularse.",
            "En Mis Postulaciones, revisar el estado de solicitudes.",
        ]),
        ("Para Autoridades", [
            "Acceder al panel con estadísticas generales.",
            "En Alumnos, filtrar por carrera, semestre y nivel de riesgo.",
            "En Reportes, consultar estadísticas, predicciones IA y exportes.",
            "Usar exportes CSV para análisis externos.",
        ]),
        ("Para Administradores", [
            "Acceder al Panel y a Gestión.",
            "Usar las pestañas de Gestión: Usuarios, Alumnos, Profesores, Académico, Operaciones y Sistema.",
            "Desde Operaciones, ejecutar evaluación masiva de riesgos.",
            "Usar Integridad de Datos para verificar consistencia.",
        ]),
    ]
    for idx, (title, items) in enumerate(sections, start=1):
        add_heading(doc, f"13.{idx}", title, 2)
        add_numbers(doc, items)


def add_technical_manual(doc: Document) -> None:
    add_heading(doc, "14.", "Manual Técnico", 1)
    add_heading(doc, "14.1", "Requisitos del Sistema", 2)
    add_bullets(doc, [".NET 8 SDK.", "SQLite en desarrollo o PostgreSQL en producción.", "Docker para despliegue en Railway.", "Git."])
    add_heading(doc, "14.2", "Instalación Local", 2)
    rows = [
        ("1", "git clone https://github.com/Xerea/RescateAcademico.git"),
        ("2", "cd RescateAcademico/RescateAcademico"),
        ("3", "dotnet restore"),
        ("4", "dotnet build"),
        ("5", "dotnet run"),
    ]
    add_table(doc, ["Paso", "Comando"], rows, [900, 8460], font_size=8)
    add_para(doc, "El sistema estará disponible en https://localhost:5001.")
    add_heading(doc, "14.3", "Estructura del Proyecto", 2)
    rows = [
        ("Controllers/", "15 controladores MVC."),
        ("Models/", "18+ modelos del dominio."),
        ("Views/", "60+ vistas Razor."),
        ("Services/", "8 servicios extraídos."),
        ("Data/", "DbContext y seeder."),
        ("Filters/", "AuditLog."),
        ("wwwroot/", "CSS, JS y librerías."),
        ("Program.cs", "Punto de entrada y configuración."),
        ("appsettings.json", "Configuración local."),
    ]
    add_table(doc, ["Ruta", "Descripción"], rows, [2600, 6760], font_size=8)
    add_heading(doc, "14.4", "Configuración de Variables de Entorno", 2)
    rows = [
        ("DATABASE_URL", "Conexión PostgreSQL (Railway)", "Producción"),
        ("OPENAI_API_KEY", "API key de OpenAI para análisis IA", "Opcional"),
        ("DEMO_ADMIN_PASSWORD", "Contraseña del admin demo", "Desarrollo"),
        ("RAILWAY_ENVIRONMENT", "Detecta entorno Railway", "Producción"),
        ("UPLOADS_PATH", "Ruta de almacenamiento de archivos", "Opcional"),
    ]
    add_table(doc, ["Variable", "Propósito", "Requerida"], rows, [2600, 5060, 1700], font_size=8)
    add_heading(doc, "14.5", "Despliegue en Railway", 2)
    add_numbers(doc, [
        "Conectar repositorio GitHub a Railway.",
        "Railway detecta automáticamente el Dockerfile y railway.toml.",
        "Agregar variables de entorno en el dashboard de Railway.",
        "El deployment se ejecuta automáticamente en cada push a main.",
        "El primer deploy tarda 8-10 minutos por el seeding de datos.",
    ])
    add_heading(doc, "14.6", "Base de Datos", 2)
    add_bullets(doc, [
        "EnsureCreated() en startup crea las tablas automáticamente.",
        "DemoDataSeeder siembra 308 estudiantes, 12 profesores y 6 carreras.",
        "Guardia de seeding: si Alumnos.Count >= 50, omite la siembra.",
        "Sin migraciones EF: cambios de esquema requieren recrear la base de datos.",
    ])


def add_implementation_plan(doc: Document) -> None:
    add_heading(doc, "15.", "Plan de Implementación", 1)
    rows = [
        ("Fase 1: Preparación", "Semana 1", "Configurar repositorio Git, CI/CD en Railway, variables de entorno, despliegue inicial y seeding."),
        ("Fase 2: Capacitación", "Semana 2", "Sesión de capacitación para 12 profesores, entrega de credenciales y guía rápida."),
        ("Fase 3: Implementación Piloto", "Semanas 3-4", "Operación con datos reales de 1 grupo piloto, feedback y ajustes menores."),
        ("Fase 4: Despliegue General", "Semana 5", "Alta de grupos y estudiantes, comunicación oficial y soporte técnico durante el primer mes."),
    ]
    add_table(doc, ["Fase", "Periodo", "Actividades"], rows, [2400, 1500, 5460], font_size=8)


def add_completion_letter(doc: Document) -> None:
    add_heading(doc, "16.", "Carta de Terminación", 1)
    add_para(doc, 'Por medio de la presente, el equipo de desarrollo del proyecto "Rescate Académico" hace constar que la plataforma web para el monitoreo de riesgo académico del IPN CECyT No. 13 "Ricardo Flores Magón" ha sido completada satisfactoriamente conforme al alcance establecido.')
    add_para(doc, "El sistema incluye autenticación por roles, dashboard institucional, perfil académico, predicción de deserción con heurística e inteligencia artificial, intervenciones de tutoría, planes de mejora, convocatorias, postulaciones, reportes, administración, bitácora de auditoría y despliegue en Railway con PostgreSQL.")
    add_para(doc, "Asimismo, se certifica que el proyecto cumple con los requisitos académicos de la asignatura de Sistema de Información y queda listo para su presentación, evaluación y operación piloto.")
    add_para(doc, "Atentamente,")
    rows = [
        ("Sergio", "Seguridad, Roles y Bitácora", "____________________________"),
        ("Sara", "Convocatorias, Postulaciones y Reportes", "____________________________"),
        ("Alejandra", "Dashboard, Perfil Académico y Estadísticas", "____________________________"),
        ("Elías", "Administración y Operación Institucional", "____________________________"),
        ("Buenfil", "Inteligencia Artificial y Predicciones", "____________________________"),
        ("Profesor asesor", "Validación académica", "____________________________"),
    ]
    add_table(doc, ["Nombre", "Responsabilidad", "Firma"], rows, [2400, 4200, 2760], font_size=8)


def add_references(doc: Document) -> None:
    add_heading(doc, "17.", "Referencias Bibliográficas", 1)
    refs = [
        "Microsoft. (2024). ASP.NET Core 8 Documentation. https://learn.microsoft.com/en-us/aspnet/core/",
        "Entity Framework Core. (2024). EF Core 8 Documentation. https://learn.microsoft.com/en-us/ef/core/",
        "OpenAI. (2024). GPT-4o-mini API Documentation. https://platform.openai.com/docs/",
        "Bootstrap. (2024). Bootstrap 5.3 Documentation. https://getbootstrap.com/docs/5.3/",
        "DataTables. (2024). DataTables 1.13 Documentation. https://datatables.net/",
        "GreenSock. (2024). GSAP 3 Documentation. https://gsap.com/docs/",
        "Railway. (2024). Railway Deployment Documentation. https://docs.railway.app/",
        "PostgreSQL. (2024). PostgreSQL 16 Documentation. https://www.postgresql.org/docs/",
        "Fowler, M. (2012). Patterns of Enterprise Application Architecture. Addison-Wesley.",
        "Instituto Politécnico Nacional. (2024). Reglamento Académico del IPN. https://www.ipn.mx/",
    ]
    add_numbers(doc, refs)


def add_maintenance(doc: Document) -> None:
    add_heading(doc, "18.", "Planificación del Mantenimiento Preventivo", 1)
    groups = [
        ("Mantenimiento Mensual", [
            "Monitoreo de logs de error en Railway.",
            "Verificación de estado de la base de datos: conexiones y espacio.",
            "Revisión de uso de API de OpenAI y costos.",
            "Backup de base de datos PostgreSQL.",
        ]),
        ("Mantenimiento Trimestral", [
            "Actualización de paquetes NuGet.",
            "Revisión de seguridad: dependencias y vulnerabilidades.",
            "Limpieza de notificaciones antiguas mayores a 30 días.",
            "Optimización de índices de base de datos.",
        ]),
        ("Mantenimiento Semestral", [
            "Renovación de API key de OpenAI.",
            "Actualización del framework de .NET 8 a .NET 9 cuando sea pertinente.",
            "Revisión de contraseñas de demo.",
            "Pruebas de rendimiento con datos reales.",
        ]),
    ]
    for idx, (title, items) in enumerate(groups, start=1):
        add_heading(doc, f"18.{idx}", title, 2)
        add_bullets(doc, items)


def add_corrective_log(doc: Document) -> None:
    add_heading(doc, "19.", "Formato de Bitácora de Mantenimiento Correctivo", 1)
    rows = [("", "", "", "", "", "", "", "") for _ in range(6)]
    add_table(doc, ["Fecha", "Hora", "Responsable", "Incidencia", "Causa", "Acción Correctiva", "Tiempo de Resolución", "Estado"], rows, [900, 750, 1450, 1450, 1200, 1850, 1200, 560], font_size=6)
    add_para(doc, "Tabla 1. Formato vacío para registro futuro de mantenimiento correctivo.", align=WD_ALIGN_PARAGRAPH.CENTER)


CONCLUSIONS = [
    ("Sergio - Seguridad, Roles y Bitácora", "Mi responsabilidad principal fue implementar y verificar la capa de seguridad del sistema. Configuré ASP.NET Core Identity con bloqueo tras 3 intentos fallidos por 20 minutos, políticas de contraseñas robustas, cookies HttpOnly, Secure y SameSite=Strict, además de rate limiting en endpoints críticos. También configuré headers de seguridad como CSP, HSTS, X-Frame-Options: DENY, X-Content-Type-Options: nosniff, Referrer-Policy y Permissions-Policy. La bitácora de auditoría registra automáticamente mutaciones críticas mediante [AuditLog], proporcionando trazabilidad completa. El mayor desafío fue equilibrar una Content Security Policy estricta con los CDNs necesarios para el frontend. Aprendí que la seguridad no se agrega al final: debe atravesar controladores, formularios, cookies, headers y servicios desde la arquitectura inicial."),
    ("Sara - Convocatorias, Postulaciones y Reportes", "Desarrollé el flujo completo de convocatorias y postulaciones académicas. Implementé el catálogo de convocatorias, el formulario de postulación, la validación de elegibilidad por promedio, semestre, carrera, cupo y carga académica, así como la gestión de estados con notificaciones automáticas. Un logro importante fue extraer la lógica a ConvocatoriaEligibilityService para aplicar las mismas reglas en GET y POST. También implementé FileStorageService con validación MIME, extensión permitida y nombres seguros, además de exportación CSV con protección contra inyección. Aprendí que centralizar reglas de negocio evita inconsistencias y vulnerabilidades."),
    ("Alejandra - Dashboard, Perfil Académico y Estadísticas", "Mi trabajo se centró en la experiencia del usuario final. Creé dashboards dinámicos por rol, el perfil académico completo del alumno, indicadores visuales, historial por periodo y el simulador What-If para proyectar escenarios de calificaciones. El hub del profesor evolucionó a un panel integral con distribución de riesgo, tarjetas de grupo, predicciones y acciones rápidas. También apoyé la migración al sistema de diseño ra-* con modo oscuro TRUE BLACK. Aprendí que una buena interfaz anticipa necesidades y reduce pasos innecesarios sin sacrificar información."),
    ("Elías - Administración y Operación Institucional", "Implementé el panel de administración con operaciones CRUD para alumnos, tutores, carreras, ciclos escolares y usuarios. Desarrollé la gestión de grupos académicos y el sistema de notificaciones con badge en tiempo real. Configuré el seeding demo con 308 estudiantes, 12 profesores, 6 carreras y calificaciones por periodos. El mayor reto fue migrar de una asignación aleatoria de tutorados a un modelo basado en grupos, más fiel a la operación institucional. Aprendí que la integridad del modelo de datos sostiene toda la funcionalidad superior."),
    ("Buenfil - Inteligencia Artificial y Predicciones", "Desarrollé el sistema de predicción de deserción con dos capas: una heurística institucional auditable y una integración con OpenAI GPT-4o-mini para análisis narrativo. La heurística pondera promedio, reprobaciones, ausencias, ETS y recursamientos; la IA genera recomendaciones estructuradas y personalizadas. Implementé persistencia, auditoría y fallback cuando OPENAI_API_KEY no está configurada. Aprendí que la IA debe aportar contexto y claridad, pero la lógica base debe seguir siendo explicable, confiable y disponible sin servicios externos."),
]


def add_conclusions(doc: Document) -> None:
    add_heading(doc, "20.", "Conclusiones por Integrante", 1)
    for idx, (title, text) in enumerate(CONCLUSIONS, start=1):
        add_heading(doc, f"20.{idx}", title, 2)
        add_para(doc, text)
    add_heading(doc, "21.", "Conclusiones Generales del Equipo", 1)
    add_para(doc, "El desarrollo de Rescate Académico representó un aprendizaje integral que abarcó desde la arquitectura de software hasta el despliegue en producción. Comprobamos que un sistema bien diseñado puede tener impacto real en la operación académica: los profesores pueden identificar estudiantes en riesgo en segundos, cuando antes requerían revisar expedientes manualmente.")
    add_para(doc, "La combinación de reglas heurísticas con inteligencia artificial generativa demostró ser un enfoque equilibrado que proporciona confiabilidad y profundidad analítica. La experiencia de desplegar en Railway reforzó la importancia de la infraestructura como código, la gestión de variables de entorno y la observabilidad.")
    add_para(doc, "Aunque el proyecto cumple con los objetivos planteados, se identifican áreas de mejora futura: adición de pruebas automatizadas, migración a paginación server-side para escalar a miles de estudiantes y notificaciones por correo electrónico real. Como equipo, aprendimos que la seguridad debe ser transversal, la lógica de negocio debe estar centralizada y la experiencia de usuario es tan importante como la funcionalidad técnica.")
    add_para(doc, "FIN DEL DOCUMENTO.", align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)


def main() -> None:
    ARTIFACTS.mkdir(exist_ok=True)
    logo = make_logo()
    er = make_er_diagram()
    classes = make_class_diagram()

    doc = Document()
    format_document(doc)
    configure_header_footer(doc.sections[0])
    add_update_fields(doc)

    add_cover(doc, logo)
    add_toc(doc)
    add_intro(doc)
    add_scope(doc)
    add_analysis(doc)
    add_backlog(doc)
    add_user_stories(doc)
    add_design(doc, er, classes)
    add_development(doc)
    add_youtube(doc)
    add_general_tests(doc)
    add_evidence(doc)
    add_user_manual(doc)
    add_technical_manual(doc)
    add_implementation_plan(doc)
    add_completion_letter(doc)
    add_references(doc)
    add_maintenance(doc)
    add_corrective_log(doc)
    add_conclusions(doc)

    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
